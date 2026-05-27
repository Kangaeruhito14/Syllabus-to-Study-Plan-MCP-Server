from __future__ import annotations

import base64
from datetime import date, datetime, timezone
from typing import Any, Literal

from fastmcp import FastMCP
from pydantic import BaseModel, Field

from syllabus_mcp.models import Assessment, AssessmentType, Confidence, CourseModel, StudyPlan, StudyPreferences, Topic
from syllabus_mcp.ocr import pdf_bytes_to_text_pages
from syllabus_mcp.extract import TextPage, extract_course_model_from_pages, extract_text_from_pdf_bytes, is_likely_syllabus
from syllabus_mcp.planner import generate_coverage_plan, generate_plan, weight_topics as weight_course_topics
from syllabus_mcp.exporters import (
    DailyNotionPushResult,
    plan_to_ics,
    push_daily_plan_to_notion,
    push_plan_to_notion,
)
from syllabus_mcp.gcal import push_plan_to_google_calendar


class ParseSyllabusInput(BaseModel):
    content_type: Literal["pdf_base64", "text", "docx_base64", "html", "url"] = Field(
        description=(
            "Input type: 'pdf_base64' (base64-encoded PDF), 'text' (plain text), "
            "'docx_base64' (base64-encoded .docx), 'html' (raw HTML string), "
            "'url' (web URL to fetch and parse)."
        )
    )
    content: str = Field(description="The content — PDF/DOCX bytes as base64, plain text, raw HTML, or a URL.")
    timezone: str | None = Field(default=None, description="Optional timezone hint.")


class ParseSyllabusOutput(BaseModel):
    course: CourseModel
    warnings: list[str] = Field(default_factory=list)
    courses_list: list[str] = Field(
        default_factory=list,
        description=(
            "Distinct course/module names found. Non-empty for multi-course program syllabi. "
            "Each entry is 'Module Name (N topics)' or the course code/title."
        ),
    )


class DetectExamDatesOutput(BaseModel):
    course: CourseModel
    warnings: list[str] = Field(default_factory=list)


class WeightTopicsInput(BaseModel):
    course: CourseModel
    boost_keywords: list[str] = Field(
        default_factory=list,
        description="Optional custom keywords to treat as importance cues.",
    )


class WeightTopicsOutput(BaseModel):
    course: CourseModel
    warnings: list[str] = Field(default_factory=list)


class GenerateStudyPlanInput(BaseModel):
    course: CourseModel
    preferences: StudyPreferences = Field(default_factory=StudyPreferences)


class GenerateStudyPlanOutput(BaseModel):
    plan: StudyPlan
    warnings: list[str] = Field(default_factory=list)


class ExportPlanInput(BaseModel):
    plan: StudyPlan
    format: Literal["json", "ics", "google_calendar", "notion"]
    target: dict[str, Any] = Field(
        default_factory=dict,
        description="Format-specific options. For integrations, include identifiers and auth hints.",
    )


class ExportPlanOutput(BaseModel):
    format: str
    result: Any
    warnings: list[str] = Field(default_factory=list)


class CourseCorrectionInput(BaseModel):
    class NewAssessmentInput(BaseModel):
        name: str
        type: AssessmentType = AssessmentType.exam
        scheduled_date: date | None = None
        weight_percent: float | None = None

    course: CourseModel
    set_course_title: str | None = None
    set_timezone: str | None = None
    set_end_date: date | None = None
    assessment_date_overrides: dict[str, date] = Field(
        default_factory=dict,
        description="Map assessment name -> corrected date.",
    )
    add_topics: list[str] = Field(default_factory=list)
    remove_topics: list[str] = Field(default_factory=list)
    remove_assessments: list[str] = Field(
        default_factory=list,
        description="Assessment names to remove (case-insensitive exact-name match).",
    )
    add_assessments: list[NewAssessmentInput] = Field(
        default_factory=list,
        description="Add or update assessments by name (manual exam/date correction flow).",
    )


class CourseCorrectionOutput(BaseModel):
    course: CourseModel
    changes_applied: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class BuildPlanReportInput(BaseModel):
    course: CourseModel
    plan: StudyPlan
    include_markdown: bool = True


class BuildPlanReportOutput(BaseModel):
    report: str


class GetRawTextInput(BaseModel):
    content_type: Literal["pdf_base64", "text", "docx_base64", "html", "url"]
    content: str = Field(description="Base64 PDF/DOCX bytes, plain text, raw HTML, or a URL.")


class GetRawTextOutput(BaseModel):
    raw_text: str
    page_count: int
    char_count: int
    is_likely_syllabus: bool
    syllabus_confidence: float
    warnings: list[str] = Field(default_factory=list)


mcp = FastMCP(
    name="Syllabus-to-Study-Plan",
    instructions=(
        "Tools to parse a syllabus (PDF/text), detect exams, weight topics, "
        "generate a spaced-repetition study schedule, and export/push to ICS, "
        "Google Calendar, or Notion. "
        "Use get_raw_text first when parse_syllabus produces low-quality results — "
        "read the raw text yourself and call apply_course_corrections to fix extraction."
    ),
)


def _decode_pdf_base64(b64: str) -> bytes:
    return base64.b64decode(b64.encode("utf-8"))


def _meaningful_chars(pages: list[TextPage]) -> int:
    """Count chars after removing scanner watermark lines and blank lines."""
    _WATERMARKS = {"scanned with", "camscanner", "created with", "adobe scan"}
    total = 0
    for p in pages:
        for line in p.text.split("\n"):
            stripped = line.strip()
            low = stripped.lower()
            if not stripped or any(w in low for w in _WATERMARKS):
                continue
            total += len(stripped)
    return total


def _html_to_text(html: str) -> str:
    """Strip HTML tags and return clean plain text."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _docx_bytes_to_text(docx_bytes: bytes) -> str:
    """Extract plain text from a .docx file."""
    import io as _io
    from docx import Document
    doc = Document(_io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _fetch_url_text(url: str) -> tuple[str, list[str]]:
    """Fetch a URL and return (plain_text, warnings)."""
    import requests as _req
    warnings: list[str] = []
    try:
        resp = _req.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0 (SyllabusBot/1.0)"})
        resp.raise_for_status()
        ctype = resp.headers.get("content-type", "")
        if "pdf" in ctype:
            warnings.append("URL returned a PDF — use content_type='pdf_base64' with the downloaded file instead.")
            return "", warnings
        text = _html_to_text(resp.text)
        if not text.strip():
            warnings.append(f"URL returned empty content. Status: {resp.status_code}")
        return text, warnings
    except Exception as exc:
        warnings.append(f"Failed to fetch URL: {exc}")
        return "", warnings


def _content_to_pages(content_type: str, content: str) -> tuple[list[TextPage], list[str]]:
    """
    Convert any supported content_type into a list of TextPage objects.
    Handles: pdf_base64, text, docx_base64, html, url.
    """
    warnings: list[str] = []

    if content_type == "pdf_base64":
        pdf_bytes = _decode_pdf_base64(content)
        pages, w = _pdf_bytes_to_pages(pdf_bytes)
        warnings.extend(w)
        return pages, warnings

    if content_type == "docx_base64":
        try:
            docx_bytes = base64.b64decode(content)
            text = _docx_bytes_to_text(docx_bytes)
            return [TextPage(page_number=1, text=text)], warnings
        except Exception as exc:
            warnings.append(f"Failed to read DOCX: {exc}")
            return [TextPage(page_number=1, text="")], warnings

    if content_type == "html":
        text = _html_to_text(content)
        return [TextPage(page_number=1, text=text)], warnings

    if content_type == "url":
        text, w = _fetch_url_text(content)
        warnings.extend(w)
        return [TextPage(page_number=1, text=text)], warnings

    # default: plain text
    return [TextPage(page_number=1, text=content.strip())], warnings


def _pdf_bytes_to_pages(pdf_bytes: bytes) -> tuple[list[TextPage], list[str]]:
    """Extract pages from PDF bytes; OCR fallback if text is sparse. Returns (pages, warnings)."""
    warnings: list[str] = []
    pages = extract_text_from_pdf_bytes(pdf_bytes)
    if _meaningful_chars(pages) < 300:
        warnings.append("PDF text extraction sparse; using OCR fallback.")
        ocr_pages = pdf_bytes_to_text_pages(pdf_bytes, max_pages=20)
        if len(pages) > 20:
            warnings.append("PDF has more than 20 pages; OCR capped at page 20 to avoid memory issues.")
        pages = [TextPage(page_number=p.page_number, text=p.text) for p in ocr_pages]
    return pages, warnings


@mcp.tool
def get_raw_text(inp: GetRawTextInput) -> GetRawTextOutput:
    """
    Extract and return the full raw text from a PDF or text input, with no parsing.

    Use this when parse_syllabus produces wrong or incomplete results.
    Read the returned raw_text yourself to identify the real course title, topics,
    and exam dates, then call apply_course_corrections to fix the CourseModel.

    Also tells you whether the input looks like a syllabus.
    """
    warnings: list[str] = []
    pages, w = _content_to_pages(inp.content_type, inp.content)
    warnings.extend(w)
    raw_text = "\n\n--- PAGE BREAK ---\n\n".join(p.text for p in pages)
    page_count = len(pages)

    likely, confidence = is_likely_syllabus(raw_text)
    if not likely:
        warnings.append(
            f"This document does not appear to be a syllabus "
            f"(confidence={confidence:.0%}). Results may be meaningless."
        )

    return GetRawTextOutput(
        raw_text=raw_text,
        page_count=page_count,
        char_count=len(raw_text),
        is_likely_syllabus=likely,
        syllabus_confidence=round(confidence, 2),
        warnings=warnings,
    )


@mcp.tool
def parse_syllabus(inp: ParseSyllabusInput) -> ParseSyllabusOutput:
    """
    Parse a syllabus (PDF or text) into a structured CourseModel.

    The extractor handles four common syllabus formats:
    - Week/Module prefix  (Week 1: Introduction, Module 2: ...)
    - Date-column table   (Jan 13   Introduction to Databases   Ch.1)
    - Course Objectives   (bullet/numbered learning outcomes section)
    - Numbered schedule   (1. Variables  2. Loops ...)

    If topics or exam dates are missing/wrong after parsing, call get_raw_text
    to read the full PDF text yourself, then fix with apply_course_corrections.
    """
    warnings: list[str] = []
    pages, w = _content_to_pages(inp.content_type, inp.content)
    warnings.extend(w)

    # Non-syllabus detection
    full_text = "\n".join(p.text for p in pages)
    likely_syl, syl_confidence = is_likely_syllabus(full_text)
    if not likely_syl:
        warnings.append(
            f"WARNING: This document does not appear to be a syllabus "
            f"(confidence={syl_confidence:.0%}). "
            "Use get_raw_text to inspect the content and verify."
        )

    course = extract_course_model_from_pages(pages, timezone=inp.timezone)

    # Quality warnings
    if not course.topics:
        warnings.append(
            "No topics detected. Call get_raw_text to read the full PDF text, "
            "then use apply_course_corrections.add_topics."
        )
    elif len(course.topics) < 3:
        warnings.append(
            f"Only {len(course.topics)} topic(s) detected — likely incomplete. "
            "Call get_raw_text to verify and fix with apply_course_corrections."
        )

    if not course.assessments:
        warnings.append(
            "No assessments detected. Use apply_course_corrections.add_assessments "
            "to add exam names and dates manually."
        )

    if any("Week-based" in c for c in course.constraints_found):
        warnings.append(
            "Week-based syllabus: no calendar dates found. "
            "Provide exam dates via apply_course_corrections or full_pipeline.manual_exam_dates."
        )

    exam_without_dates = [a for a in course.assessments if a.type == AssessmentType.exam and not a.scheduled_date]
    if exam_without_dates:
        names = ", ".join(a.name for a in exam_without_dates[:3])
        warnings.append(f"Exams detected without dates: {names}. Add dates via apply_course_corrections.")

    # Build courses_list by grouping topics on their module field
    from collections import Counter as _Counter
    module_counts: dict[str, int] = {}
    for t in course.topics:
        if t.module:
            module_counts[t.module] = module_counts.get(t.module, 0) + 1
    courses_list: list[str] = [
        f"{mod} ({cnt} topic{'s' if cnt != 1 else ''})"
        for mod, cnt in module_counts.items()
    ]

    return ParseSyllabusOutput(course=course, warnings=warnings, courses_list=courses_list)


@mcp.tool
def detect_exam_dates(course: CourseModel) -> DetectExamDatesOutput:
    """Refine assessments and exam dates with confidence and assumptions."""
    warnings: list[str] = []
    blockers = ("assignment must", "blackboard", "attendance", "office hours")
    exams = []
    for a in course.assessments:
        name_low = a.name.lower()
        if any(b in name_low for b in blockers):
            continue
        if a.type == AssessmentType.exam or any(k in name_low for k in ("midterm", "final", "exam", "test")):
            if a.scheduled_date is None:
                warnings.append(f"Missing date for likely exam: '{a.name}'.")
            exams.append(a)

    if exams and any(e.scheduled_date for e in exams):
        course.end_date = max([e.scheduled_date for e in exams if e.scheduled_date])
        course.confidences["end_date"] = Confidence(value=0.55, reason="Derived from detected exam assessments")
    elif course.end_date is None:
        warnings.append("No exam dates detected. Provide overrides via apply_course_corrections.")

    return DetectExamDatesOutput(course=course, warnings=warnings)


@mcp.tool
def weight_topics(inp: WeightTopicsInput) -> WeightTopicsOutput:
    """Assign weight scores to topics with rationale."""
    course = weight_course_topics(inp.course, boost_keywords=inp.boost_keywords)
    return WeightTopicsOutput(course=course, warnings=[])


@mcp.tool
def generate_study_plan(inp: GenerateStudyPlanInput) -> GenerateStudyPlanOutput:
    """Generate a day-by-day schedule with spaced repetition based on preferences."""
    plan = generate_plan(inp.course, inp.preferences)
    plan.meta["generated_at"] = datetime.now(timezone.utc).isoformat()
    warnings: list[str] = []
    if not inp.course.topics:
        warnings.append("No topics found; plan will contain practice/buffer only.")
    if not any(a.scheduled_date for a in inp.course.assessments) and not inp.course.end_date:
        warnings.append("No exam/end date detected; plan uses a 30-day default window.")
    return GenerateStudyPlanOutput(plan=plan, warnings=warnings)


@mcp.tool
def export_plan(inp: ExportPlanInput) -> ExportPlanOutput:
    """Export a plan to JSON/ICS or push to Google Calendar/Notion."""
    if inp.format == "json":
        return ExportPlanOutput(format="json", result=inp.plan.model_dump(), warnings=[])

    if inp.format == "ics":
        ics_text = plan_to_ics(inp.plan)
        return ExportPlanOutput(format="ics", result={"ics": ics_text}, warnings=[])

    if inp.format == "notion":
        token = inp.target.get("notion_token")
        database_id = inp.target.get("database_id")
        if not token or not database_id:
            return ExportPlanOutput(
                format="notion",
                result={"status": "missing_credentials"},
                warnings=["Provide target.notion_token and target.database_id."],
            )
        res = push_plan_to_notion(
            inp.plan,
            notion_token=token,
            database_id=database_id,
            prune_stale=bool(inp.target.get("prune_stale", True)),
        )
        return ExportPlanOutput(format="notion", result=res.__dict__, warnings=[])

    if inp.format == "google_calendar":
        access_token = inp.target.get("access_token")
        if not access_token:
            return ExportPlanOutput(
                format="google_calendar",
                result={"status": "missing_credentials"},
                warnings=["Provide target.access_token (and optionally refresh_token, client_id, client_secret)."],
            )
        res = push_plan_to_google_calendar(
            inp.plan,
            access_token=access_token,
            refresh_token=inp.target.get("refresh_token"),
            client_id=inp.target.get("client_id"),
            client_secret=inp.target.get("client_secret"),
            calendar_id=inp.target.get("calendar_id", "primary"),
            prune_stale=bool(inp.target.get("prune_stale", True)),
        )
        return ExportPlanOutput(format="google_calendar", result=res.__dict__, warnings=[])

    return ExportPlanOutput(format=inp.format, result={"status": "unknown_format"}, warnings=[])


@mcp.tool
def apply_course_corrections(inp: CourseCorrectionInput) -> CourseCorrectionOutput:
    """
    Apply user-specified corrections (title, dates, topics) to a parsed course model.
    """
    course = inp.course.model_copy(deep=True)
    changes: list[str] = []
    warnings: list[str] = []

    if inp.set_course_title:
        course.course_title = inp.set_course_title.strip()
        changes.append(f"course_title='{course.course_title}'")
    if inp.set_timezone:
        course.timezone = inp.set_timezone.strip()
        changes.append(f"timezone='{course.timezone}'")
    if inp.set_end_date:
        course.end_date = inp.set_end_date
        changes.append(f"end_date='{inp.set_end_date.isoformat()}'")

    if inp.assessment_date_overrides:
        lowered = {k.strip().lower(): v for k, v in inp.assessment_date_overrides.items()}
        matched = 0
        for a in course.assessments:
            key = a.name.strip().lower()
            if key in lowered:
                a.scheduled_date = lowered[key]
                if a.confidence:
                    a.confidence = Confidence(value=0.95, reason="User-corrected date")
                matched += 1
        if matched == 0:
            warnings.append("No assessment names matched assessment_date_overrides.")
        else:
            changes.append(f"assessment_dates_updated={matched}")

    if inp.add_assessments:
        existing_by_name = {a.name.strip().lower(): a for a in course.assessments}
        added = 0
        updated = 0
        for item in inp.add_assessments:
            key = item.name.strip().lower()
            if not key:
                continue
            if key in existing_by_name:
                a = existing_by_name[key]
                a.type = item.type
                if item.scheduled_date is not None:
                    a.scheduled_date = item.scheduled_date
                if item.weight_percent is not None:
                    a.weight_percent = item.weight_percent
                a.confidence = Confidence(value=0.95, reason="User-added assessment details")
                updated += 1
                continue

            course.assessments.append(
                Assessment(
                    name=item.name.strip(),
                    type=item.type,
                    scheduled_date=item.scheduled_date,
                    weight_percent=item.weight_percent,
                    confidence=Confidence(value=0.98, reason="User-added assessment"),
                    source={"page": None, "line": "manual_add_assessment"},
                )
            )
            existing_by_name[key] = course.assessments[-1]
            added += 1

        if added:
            changes.append(f"assessments_added={added}")
        if updated:
            changes.append(f"assessments_updated={updated}")

    remove_assessment_set = {a.strip().lower() for a in inp.remove_assessments if a.strip()}
    if remove_assessment_set:
        before = len(course.assessments)
        course.assessments = [
            a for a in course.assessments if a.name.strip().lower() not in remove_assessment_set
        ]
        removed = before - len(course.assessments)
        if removed:
            changes.append(f"assessments_removed={removed}")
        else:
            warnings.append("No assessment names matched remove_assessments.")

    remove_set = {t.strip().lower() for t in inp.remove_topics if t.strip()}
    if remove_set:
        before = len(course.topics)
        course.topics = [t for t in course.topics if t.title.strip().lower() not in remove_set]
        removed = before - len(course.topics)
        changes.append(f"topics_removed={removed}")

    existing = {t.title.strip().lower() for t in course.topics}
    added = 0
    for raw in inp.add_topics:
        title = raw.strip()
        if not title:
            continue
        if title.lower() in existing:
            continue
        course.topics.append(Topic(title=title))
        existing.add(title.lower())
        added += 1
    if added:
        changes.append(f"topics_added={added}")

    dated = [a.scheduled_date for a in course.assessments if a.scheduled_date]
    if dated:
        course.end_date = max(dated)

    return CourseCorrectionOutput(course=course, changes_applied=changes, warnings=warnings)


@mcp.tool
def build_plan_report(inp: BuildPlanReportInput) -> BuildPlanReportOutput:
    """
    Build a polished natural-language report for users.

    Includes: exam countdown, weekly session summary, priority topics,
    first sessions preview, and actionable adjustment tips.
    """
    course = inp.course
    plan = inp.plan
    today = date.today()
    md = inp.include_markdown

    exams = [a for a in course.assessments if a.type == AssessmentType.exam]
    dated_exams = sorted([a for a in exams if a.scheduled_date], key=lambda x: x.scheduled_date)
    total_hours = round(sum(s.estimated_minutes for s in plan.sessions) / 60, 1)
    total_days = (plan.end_date - plan.start_date).days + 1
    learn_count = sum(1 for s in plan.sessions if s.session_type.value == "learn")
    review_count = sum(1 for s in plan.sessions if s.session_type.value == "review")

    lines: list[str] = []

    # ── Header ────────────────────────────────────────────────────────────────
    title = course.course_title or "Your Course"
    if md:
        lines += [f"# Study Plan — {title}", ""]

    # ── Snapshot ──────────────────────────────────────────────────────────────
    if md:
        lines += ["## At a Glance", ""]
    lines += [
        f"- Course: {title}",
        f"- Study window: {plan.start_date.isoformat()} → {plan.end_date.isoformat()} ({total_days} days)",
        f"- Total sessions: {len(plan.sessions)}  ({learn_count} learn · {review_count} review)",
        f"- Estimated total study time: {total_hours} hours",
        f"- Topics in plan: {len(course.topics)}",
        f"- Assessments tracked: {len(course.assessments)}",
    ]

    # ── Exam Countdown ────────────────────────────────────────────────────────
    if md:
        lines += ["", "## Exam Countdown", ""]
    else:
        lines.append("")
    if dated_exams:
        for a in dated_exams:
            days_left = (a.scheduled_date - today).days
            if days_left < 0:
                tag = "(past)"
            elif days_left == 0:
                tag = "TODAY"
            elif days_left == 1:
                tag = "tomorrow"
            else:
                tag = f"{days_left} days away"
            lines.append(f"- {a.name}: {a.scheduled_date.isoformat()}  [{tag}]")
    else:
        lines.append(
            "- No exam dates detected yet. "
            "Use `apply_course_corrections` or `full_pipeline.manual_exam_dates` to add them."
        )

    # ── Weekly Summary ────────────────────────────────────────────────────────
    if md:
        lines += ["", "## Weekly Study Summary", ""]
    else:
        lines.append("")

    from collections import defaultdict
    week_map: dict[str, int] = defaultdict(int)
    for s in plan.sessions:
        week_num = s.session_date.isocalendar()[1]
        year = s.session_date.year
        key = f"{year}-W{week_num:02d}"
        week_map[key] += s.estimated_minutes

    shown = 0
    for week_key in sorted(week_map.keys())[:6]:
        mins = week_map[week_key]
        hrs = round(mins / 60, 1)
        lines.append(f"- {week_key}: {hrs} hrs ({mins} min)")
        shown += 1
    if len(week_map) > shown:
        lines.append(f"- … and {len(week_map) - shown} more weeks")

    # ── Priority Topics ───────────────────────────────────────────────────────
    if md:
        lines += ["", "## Priority Topics (highest weight first)", ""]
    else:
        lines.append("")

    top_topics = sorted(course.topics, key=lambda t: t.weight_score or 1.0, reverse=True)[:8]
    if top_topics:
        for t in top_topics:
            score = t.weight_score or 1.0
            bar = "█" * min(int(score), 5) + "░" * max(0, 5 - int(score))
            lines.append(f"- {t.title}  [{bar}] {score:.1f}")
    else:
        lines.append("- No topics detected. Use `apply_course_corrections.add_topics` to add them.")

    # ── First Sessions ────────────────────────────────────────────────────────
    if md:
        lines += ["", "## Your First 7 Sessions", ""]
    else:
        lines.append("")

    for s in plan.sessions[:7]:
        lines.append(
            f"- {s.session_date.isoformat()} [{s.session_type.value:8s}] "
            f"{s.topic_title} ({s.estimated_minutes} min)"
        )

    # ── How to Adjust ─────────────────────────────────────────────────────────
    if md:
        lines += ["", "## How to Adjust This Plan", ""]
    else:
        lines.append("")

    lines += [
        "- Wrong exam dates? → call `apply_course_corrections` with `assessment_date_overrides`, then `generate_study_plan`.",
        "- Missing topics?   → call `apply_course_corrections` with `add_topics`, then `weight_topics` + `generate_study_plan`.",
        "- Change study hours or days off? → call `generate_study_plan` again with updated `preferences`.",
        "- Re-sync calendar? → call `export_plan` with `format='ics'` (or `google_calendar`/`notion`) after regenerating.",
    ]

    return BuildPlanReportOutput(report="\n".join(lines))


class SetupNotionDatabaseInput(BaseModel):
    notion_token: str = Field(description="Your Notion integration token (starts with 'secret_...').")
    parent_page_id: str = Field(
        description=(
            "The Notion page ID where the database will be created. "
            "Open the page in Notion, copy the URL — the ID is the last 32-char hex after the last '/'."
        )
    )
    database_title: str = Field(default="Study Plan", description="Title for the new Notion database.")


class SetupNotionDatabaseOutput(BaseModel):
    database_id: str
    database_url: str
    message: str
    warnings: list[str] = Field(default_factory=list)


@mcp.tool
def setup_notion_database(inp: SetupNotionDatabaseInput) -> SetupNotionDatabaseOutput:
    """
    Create the required Notion database for study plan export.

    Creates a database with all required properties:
    Name, Date, Type, Minutes, Key, Rationale, PlanKey.

    Run this once, then use the returned database_id in export_plan or full_pipeline.
    """
    import requests as _req

    headers = {
        "Authorization": f"Bearer {inp.notion_token}",
        "Notion-Version": "2022-06-28",
        "Content-Type": "application/json",
    }
    payload = {
        "parent": {"type": "page_id", "page_id": inp.parent_page_id},
        "title": [{"type": "text", "text": {"content": inp.database_title}}],
        "properties": {
            "Name": {"title": {}},
            "Date": {"date": {}},
            "Type": {
                "select": {
                    "options": [
                        {"name": "learn", "color": "blue"},
                        {"name": "review", "color": "green"},
                        {"name": "practice", "color": "yellow"},
                        {"name": "mock_exam", "color": "red"},
                        {"name": "buffer", "color": "gray"},
                    ]
                }
            },
            "Minutes": {"number": {"format": "number"}},
            "Key": {"rich_text": {}},
            "PlanKey": {"rich_text": {}},
            "Rationale": {"rich_text": {}},
        },
    }

    try:
        resp = _req.post(
            "https://api.notion.com/v1/databases",
            headers=headers,
            json=payload,
            timeout=30,
        )
        resp.raise_for_status()
        db = resp.json()
        db_id = db["id"]
        db_url = db.get("url", f"https://notion.so/{db_id.replace('-', '')}")
        return SetupNotionDatabaseOutput(
            database_id=db_id,
            database_url=db_url,
            message=(
                f"Database '{inp.database_title}' created successfully. "
                f"Use database_id='{db_id}' in export_plan or full_pipeline."
            ),
            warnings=[],
        )
    except Exception as exc:
        return SetupNotionDatabaseOutput(
            database_id="",
            database_url="",
            message="Failed to create database. See warnings.",
            warnings=[str(exc)],
        )


class SetupDailyNotionDatabaseInput(BaseModel):
    notion_token: str = Field(description="Your Notion integration token (starts with 'secret_...').")
    parent_page_id: str = Field(
        description=(
            "The Notion page ID where the database will be created. "
            "Open the page in Notion, copy the URL — the ID is the 32-char hex after the last '/'."
        )
    )
    database_title: str = Field(default="Daily Study Plan", description="Title for the new database.")


class SetupDailyNotionDatabaseOutput(BaseModel):
    database_id: str
    database_url: str
    message: str
    warnings: list[str] = Field(default_factory=list)


@mcp.tool
def setup_daily_notion_database(inp: SetupDailyNotionDatabaseInput) -> SetupDailyNotionDatabaseOutput:
    """
    Create the Notion database for the DAILY study plan (one row per day).

    Schema: Name | Date | Day | Topics | Details | Total Minutes | Done (checkbox)

    - Name:          "2026-06-02 — Monday"
    - Date:          the calendar date
    - Day:           "Monday, 02 Jun 2026"
    - Topics:        all topic titles for that day (newline-separated)
    - Details:       per-session breakdown "[learn] Topic — 60 min"
    - Total Minutes: total study time that day
    - Done:          checkbox to mark the day complete
    - Key / PlanKey: internal idempotency keys (do not edit)

    Run this once, then use the returned database_id as notion_daily_database_id in full_pipeline.
    """
    import requests as _req

    headers = {
        "Authorization": f"Bearer {inp.notion_token}",
        "Notion-Version": "2022-06-28",
        "Content-Type": "application/json",
    }
    payload = {
        "parent": {"type": "page_id", "page_id": inp.parent_page_id},
        "title": [{"type": "text", "text": {"content": inp.database_title}}],
        "properties": {
            "Name": {"title": {}},
            "Date": {"date": {}},
            "Day": {"rich_text": {}},
            "Topics": {"rich_text": {}},
            "Details": {"rich_text": {}},
            "Total Minutes": {"number": {"format": "number"}},
            "Done": {"checkbox": {}},
            "Key": {"rich_text": {}},
            "PlanKey": {"rich_text": {}},
        },
    }
    try:
        resp = _req.post(
            "https://api.notion.com/v1/databases",
            headers=headers,
            json=payload,
            timeout=30,
        )
        resp.raise_for_status()
        db = resp.json()
        db_id = db["id"]
        db_url = db.get("url", f"https://notion.so/{db_id.replace('-', '')}")
        return SetupDailyNotionDatabaseOutput(
            database_id=db_id,
            database_url=db_url,
            message=(
                f"Daily study plan database '{inp.database_title}' created. "
                f"Use database_id='{db_id}' as notion_daily_database_id in full_pipeline."
            ),
            warnings=[],
        )
    except Exception as exc:
        return SetupDailyNotionDatabaseOutput(
            database_id="",
            database_url="",
            message="Failed to create database. See warnings.",
            warnings=[str(exc)],
        )


class FullPipelineInput(BaseModel):
    content_type: Literal["pdf_base64", "text", "docx_base64", "html", "url"] = Field(
        description=(
            "Input type: 'pdf_base64', 'text', 'docx_base64', 'html', or 'url'."
        )
    )
    content: str = Field(description="Base64 PDF/DOCX bytes, plain text, raw HTML, or a URL.")
    timezone: str = Field(default="UTC", description="Your timezone, e.g. 'Asia/Kolkata'.")
    hours_per_day: float = Field(default=1.5, description="Study hours available per day.")
    session_minutes: int = Field(
        default=60,
        ge=15,
        le=180,
        description="Minutes of focused study per session (per day in coverage mode). Default 60.",
    )
    days_off: list[Literal["mon", "tue", "wed", "thu", "fri", "sat", "sun"]] = Field(
        default_factory=list,
        description="Days you cannot study, e.g. ['fri', 'sat'].",
    )
    course_start_date: date | None = Field(
        default=None,
        description="When to start studying. Defaults to today.",
    )

    # ── Plan mode ─────────────────────────────────────────────────────────────
    plan_mode: Literal["spaced_repetition", "coverage"] = Field(
        default="spaced_repetition",
        description=(
            "'spaced_repetition': classic SR schedule anchored to exam dates. "
            "'coverage': daily sequential coverage — one topic per day, spread evenly "
            "from course_start_date to study_end_date. No exam date needed."
        ),
    )
    study_end_date: str | None = Field(
        default=None,
        description=(
            "End date for coverage mode (ISO format 'YYYY-MM-DD'). "
            "E.g. 4 months from today for a semester plan. "
            "Defaults to course_start_date + 120 days if omitted."
        ),
    )
    tutorial_dates: list[dict] | None = Field(
        default=None,
        description=(
            "Upcoming tutorial/test dates. Each entry must be a dict with keys: "
            "'date' (ISO date), 'topic_hint' (topic name), 'prep_days' (int, default 3). "
            "Example: [{\"date\": \"2026-07-20\", \"topic_hint\": \"Data Structures\", \"prep_days\": 3}]. "
            "Tutorial days are blocked on the calendar; the N study days before each tutorial "
            "become focused prep days for that topic only."
        ),
    )

    # ── Exam overrides (spaced_repetition mode) ────────────────────────────────
    manual_exam_dates: dict[str, str] = Field(
        default_factory=dict,
        description=(
            "Override or add exam dates. Key = exam name, value = ISO date 'YYYY-MM-DD'. "
            "Use when the syllabus has no explicit dates (e.g. week-based schedules)."
        ),
    )
    boost_keywords: list[str] = Field(
        default_factory=list,
        description="Extra keywords to boost topic importance (e.g. ['networking', 'security']).",
    )

    # ── Scheduling behaviour ──────────────────────────────────────────────────
    study_start_hour: int = Field(
        default=20,
        ge=0,
        le=23,
        description=(
            "Hour of day (0-23) when study sessions start on the calendar. "
            "Default 20 = 8 PM, matching the common student evening study window (6-10 PM). "
            "Days with both a review + learn session stack them: review at study_start_hour, "
            "then learn 5 min later."
        ),
    )
    next_day_review: bool = Field(
        default=True,
        description=(
            "Coverage mode: automatically schedule a short review the day AFTER each learn session. "
            "Day N = learn new topic; Day N+1 = review yesterday's topic (review_minutes) "
            "then learn the next new topic. This mirrors how memory consolidation works and "
            "is far more effective than batching all reviews at the end."
        ),
    )
    review_minutes: int = Field(
        default=25,
        ge=10,
        le=60,
        description="Duration of next-day review sessions in coverage mode (default 25 min).",
    )

    # ── Export ────────────────────────────────────────────────────────────────
    export_format: Literal["ics", "json", "notion", "notion_daily", "google_calendar"] = Field(
        default="ics",
        description=(
            "Output format: "
            "'ics' — importable .ics calendar file; "
            "'json' — raw plan data; "
            "'notion' — push to Notion (one row per session, requires notion_token + notion_database_id); "
            "'notion_daily' — push to Notion as a daily plan (one row per day with all topics listed, "
            "requires notion_token + notion_daily_database_id — run setup_daily_notion_database first); "
            "'google_calendar' — push to Google Calendar (requires gcal_access_token)."
        ),
    )

    # ── Notion credentials ─────────────────────────────────────────────────────
    notion_token: str | None = Field(
        default=None,
        description="Notion integration token (starts with 'secret_...'). Required for notion exports.",
    )
    notion_database_id: str | None = Field(
        default=None,
        description="Notion database ID for 'notion' export (session-per-row). Use setup_notion_database to create it.",
    )
    notion_daily_database_id: str | None = Field(
        default=None,
        description=(
            "Notion database ID for 'notion_daily' export (day-per-row). "
            "Use setup_daily_notion_database to create it."
        ),
    )

    # ── Google Calendar credentials ────────────────────────────────────────────
    gcal_access_token: str | None = Field(
        default=None,
        description="Google OAuth2 access token. Required for 'google_calendar' export.",
    )
    gcal_refresh_token: str | None = Field(default=None, description="Google OAuth2 refresh token.")
    gcal_client_id: str | None = Field(default=None, description="Google OAuth2 client ID.")
    gcal_client_secret: str | None = Field(default=None, description="Google OAuth2 client secret.")
    gcal_calendar_id: str = Field(default="primary", description="Google Calendar ID to push events to.")


class FullPipelineOutput(BaseModel):
    course_title: str | None
    topics_found: int
    assessments_found: int
    sessions_total: int
    plan_start: str
    plan_end: str
    plan_mode: str = "spaced_repetition"
    tutorial_days_scheduled: int = 0
    prep_days_scheduled: int = 0
    report: str
    export_format: str
    export_data: Any
    warnings: list[str] = Field(default_factory=list)
    is_syllabus: bool = True
    courses_list: list[str] = Field(
        default_factory=list,
        description="Distinct courses/modules found (non-empty for multi-course program syllabi).",
    )
    raw_text_preview: str = Field(
        default="",
        description=(
            "First 1500 chars of extracted PDF text. "
            "If topics/title look wrong, read this and call apply_course_corrections to fix."
        ),
    )


@mcp.tool
def full_pipeline(inp: FullPipelineInput) -> FullPipelineOutput:
    """
    One-call end-to-end pipeline: upload a syllabus and get a complete study plan.

    ── PLAN MODES ────────────────────────────────────────────────────────────
    plan_mode='coverage' (recommended for semester planning):
      No exam date needed. Topics are assigned sequentially one per day.
      next_day_review=True (default): Day N = learn new topic (session_minutes),
      Day N+1 = review yesterday (review_minutes, shown BEFORE that day's new topic).
      This mirrors how memory consolidation works — not "learn everything, revise later".
      tutorial_dates: block tutorial days + N prep days before each one automatically.

    plan_mode='spaced_repetition':
      Classic SR schedule anchored to exam dates.
      Supply exam dates via manual_exam_dates when the syllabus has none.

    ── STUDY TIMING ──────────────────────────────────────────────────────────
    study_start_hour (default 20 = 8 PM): all calendar events start at this hour.
    On days with both a review + new topic, events stack:
      e.g. 20:00-20:25 [review] Yesterday's topic → 20:30-21:30 [learn] New topic.

    ── DYNAMIC UPDATES ───────────────────────────────────────────────────────
    Re-calling full_pipeline with ANY changed parameter (different days_off, new
    tutorial_dates, different study_start_hour, etc.) will automatically:
      - Update changed sessions in Notion and Google Calendar
      - Delete sessions that no longer exist in the new plan
      - Create new sessions
    This is fully idempotent — no duplicate events are created.
    Just call full_pipeline again with the updated params whenever the plan changes.

    ── EXPORTS ───────────────────────────────────────────────────────────────
    'ics'            → importable .ics calendar file (returned in export_data.ics)
    'json'           → raw plan data
    'notion'         → push to Notion (one row/session) — needs notion_token + notion_database_id
    'notion_daily'   → push to Notion (one row/day, with Topics + Details columns)
                       run setup_daily_notion_database once to create the schema
    'google_calendar'→ push to Google Calendar — needs gcal_access_token

    If topics or title look wrong, check raw_text_preview and call apply_course_corrections.
    """
    from datetime import timedelta as _td
    warnings: list[str] = []

    # 1) Parse
    parse_out = parse_syllabus(ParseSyllabusInput(
        content_type=inp.content_type,
        content=inp.content,
        timezone=inp.timezone,
    ))
    warnings.extend(parse_out.warnings)

    # Non-syllabus early signal
    raw_text = parse_out.course.extracted_text_summary or ""
    is_syl, syl_conf = is_likely_syllabus(raw_text)
    if not is_syl:
        warnings.append(
            f"This document may not be a syllabus (confidence={syl_conf:.0%}). "
            "Check raw_text_preview to verify."
        )

    # 2) Detect exam dates
    detect_out = detect_exam_dates(parse_out.course)
    warnings.extend(detect_out.warnings)

    # 3) Apply manual exam date overrides / additions (if provided)
    course = detect_out.course
    if inp.manual_exam_dates:
        parsed_overrides: dict[str, date] = {}
        new_assessments: list[CourseCorrectionInput.NewAssessmentInput] = []
        existing_names = {a.name.strip().lower() for a in course.assessments}
        for name, raw_date in inp.manual_exam_dates.items():
            try:
                from dateutil import parser as dp
                d = dp.parse(raw_date).date()
            except Exception:
                warnings.append(f"Could not parse date '{raw_date}' for '{name}'; skipped.")
                continue
            if name.strip().lower() in existing_names:
                parsed_overrides[name] = d
            else:
                new_assessments.append(
                    CourseCorrectionInput.NewAssessmentInput(name=name, type=AssessmentType.exam, scheduled_date=d)
                )
        corr_out = apply_course_corrections(CourseCorrectionInput(
            course=course,
            assessment_date_overrides=parsed_overrides,
            add_assessments=new_assessments,
        ))
        warnings.extend(corr_out.warnings)
        course = corr_out.course

    # 4) Weight topics
    weight_out = weight_topics(WeightTopicsInput(course=course, boost_keywords=inp.boost_keywords))
    course = weight_out.course

    # 5) Build study preferences
    prefs = StudyPreferences(
        course_start_date=inp.course_start_date,
        timezone=inp.timezone,
        hours_per_day=inp.hours_per_day,
        session_minutes=inp.session_minutes,
        days_off=inp.days_off,
    )

    # 6) Generate plan (mode-aware)
    tutorial_days_scheduled = 0
    prep_days_scheduled = 0

    if inp.plan_mode == "coverage":
        # Resolve study_end_date
        start = inp.course_start_date or date.today()
        if inp.study_end_date:
            try:
                from dateutil import parser as _dp2
                end_date = _dp2.parse(inp.study_end_date).date()
            except Exception:
                warnings.append(f"Could not parse study_end_date '{inp.study_end_date}'; defaulting to start + 120 days.")
                end_date = start + _td(days=120)
        else:
            end_date = start + _td(days=120)
            warnings.append(
                "study_end_date not provided for coverage mode; defaulting to 4 months (120 days) from start."
            )

        if not course.topics:
            warnings.append(
                "No topics detected — coverage plan will contain revision sessions only. "
                "Call get_raw_text to inspect the syllabus and use apply_course_corrections to add topics."
            )

        plan = generate_coverage_plan(
            course, prefs,
            study_end_date=end_date,
            tutorial_dates=inp.tutorial_dates,
            next_day_review=inp.next_day_review,
            review_minutes=inp.review_minutes,
        )
        plan.meta["generated_at"] = datetime.now(timezone.utc).isoformat()

        from syllabus_mcp.models import SessionType as _ST
        tutorial_days_scheduled = sum(1 for s in plan.sessions if s.session_type == _ST.tutorial)
        prep_days_scheduled = sum(1 for s in plan.sessions if s.session_type == _ST.tutorial_prep)

        if inp.tutorial_dates and tutorial_days_scheduled == 0:
            warnings.append(
                "tutorial_dates were provided but no tutorial sessions were scheduled — "
                "check that dates fall within the study window."
            )

    else:  # spaced_repetition
        plan_out = generate_study_plan(GenerateStudyPlanInput(course=course, preferences=prefs))
        warnings.extend(plan_out.warnings)
        plan = plan_out.plan
        plan.meta["generated_at"] = datetime.now(timezone.utc).isoformat()

    # 7) Build human-readable report
    report_out = build_plan_report(BuildPlanReportInput(course=course, plan=plan, include_markdown=True))

    # 8) Export
    export_format = inp.export_format
    export_data: Any = None
    export_warnings: list[str] = []

    if export_format == "json":
        export_data = plan.model_dump()

    elif export_format == "ics":
        export_data = {"ics": plan_to_ics(plan, study_start_hour=inp.study_start_hour)}

    elif export_format == "notion":
        if not inp.notion_token:
            export_warnings.append("Provide notion_token for Notion export.")
        elif not inp.notion_database_id:
            export_warnings.append(
                "Provide notion_database_id for Notion export. "
                "Run setup_notion_database to create the database first."
            )
        else:
            try:
                res = push_plan_to_notion(
                    plan,
                    notion_token=inp.notion_token,
                    database_id=inp.notion_database_id,
                )
                export_data = {"created": res.created, "updated": res.updated, "database_id": res.database_id}
            except Exception as exc:
                export_warnings.append(f"Notion export failed: {exc}")

    elif export_format == "notion_daily":
        if not inp.notion_token:
            export_warnings.append("Provide notion_token for Notion daily export.")
        elif not inp.notion_daily_database_id:
            export_warnings.append(
                "Provide notion_daily_database_id for daily Notion export. "
                "Run setup_daily_notion_database to create the database first."
            )
        else:
            try:
                res = push_daily_plan_to_notion(
                    plan,
                    notion_token=inp.notion_token,
                    database_id=inp.notion_daily_database_id,
                )
                export_data = {
                    "created": res.created,
                    "updated": res.updated,
                    "days_total": res.days_total,
                    "database_id": res.database_id,
                }
            except Exception as exc:
                export_warnings.append(f"Notion daily export failed: {exc}")

    elif export_format == "google_calendar":
        if not inp.gcal_access_token:
            export_warnings.append("Provide gcal_access_token for Google Calendar export.")
        else:
            try:
                gcal_res = push_plan_to_google_calendar(
                    plan,
                    access_token=inp.gcal_access_token,
                    refresh_token=inp.gcal_refresh_token,
                    client_id=inp.gcal_client_id,
                    client_secret=inp.gcal_client_secret,
                    calendar_id=inp.gcal_calendar_id,
                    study_start_hour=inp.study_start_hour,
                )
                export_data = gcal_res.__dict__
            except Exception as exc:
                export_warnings.append(f"Google Calendar export failed: {exc}")

    warnings.extend(export_warnings)

    return FullPipelineOutput(
        course_title=course.course_title,
        topics_found=len(course.topics),
        assessments_found=len(course.assessments),
        sessions_total=len(plan.sessions),
        plan_start=plan.start_date.isoformat(),
        plan_end=plan.end_date.isoformat(),
        plan_mode=inp.plan_mode,
        tutorial_days_scheduled=tutorial_days_scheduled,
        prep_days_scheduled=prep_days_scheduled,
        report=report_out.report,
        export_format=export_format,
        export_data=export_data,
        warnings=warnings,
        is_syllabus=is_syl,
        courses_list=parse_out.courses_list,
        raw_text_preview=(course.extracted_text_summary or "")[:1500],
    )


def main() -> None:
    mcp.run()


if __name__ == "__main__":
    main()

